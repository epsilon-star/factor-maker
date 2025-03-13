import docx
from docx.enum.table import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import docxtpl
from docx2pdf import convert

from fanumber import convert_to_words

# doc = docx.Document("factors/template.docx")

def fa_number(number):
    persian_digits = {
        '0': '۰', '1': '۱', '2': '۲', '3': '۳', '4': '۴', 
        '5': '۵', '6': '۶', '7': '۷', '8': '۸', '9': '۹'
    }
    return ''.join(persian_digits[digit] if digit in persian_digits else digit for digit in str(number))

def formatNumber(number):
    output = str(fa_number(number))[::-1]  # Reverse the string
    grouped = [output[i:i+3] for i in range(0, len(output), 3)]  # Split into groups of 3
    output_with_commas = "،".join(grouped)[::-1]  # Join with commas and reverse back
    return output_with_commas

# doc.tables[0].cell(1,1).text = "بلوک " + formatNumber(1523642)
# doc.tables[0].style.font.name = "Vazir"

# doc.save("test.docx")


class FactorMaker:
    def createFactor(template,output,datas,items,metas):
        if template:
            doc = docx.Document(template)

            itemtable = doc.tables[2]

            total_price = 0
            total_items = 0
            price_row = []
            row_gap = 2

            for x in range(len(items)):
                itemtable.add_row()

                cell = itemtable.cell(x+row_gap,4)
                cell.paragraphs[0].text = items[x][0]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.paragraphs[0].style.font.size = Pt(8)

                cell = itemtable.cell(x+row_gap,3)
                cell.paragraphs[0].text = items[x][1]
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.paragraphs[0].style.font.size = Pt(8)

                cell = itemtable.cell(x+row_gap,2)
                cell.paragraphs[0].text = formatNumber(str(items[x][-2]))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.paragraphs[0].style.font.size = Pt(8)

                cell = itemtable.cell(x+row_gap,1)
                cell.paragraphs[0].text = formatNumber(str(items[x][-1]))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.paragraphs[0].style.font.size = Pt(8)

                cell = itemtable.cell(x+row_gap,0)
                cell.paragraphs[0].text = formatNumber(str(items[x][-1] * items[x][-2]))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cell.paragraphs[0].style.font.size = Pt(8)

                total_price += items[x][-1] * items[x][-2]
                total_items += items[x][-2]
                price_row.append(items[x][-1] * items[x][-2])

            itemtable.add_row()
            cell = itemtable.cell(len(items) + row_gap,4)
            cell.merge(itemtable.cell(len(items) + row_gap,3))
            cell = itemtable.cell(len(items) + row_gap,3)
            cell.paragraphs[0].text = "جمع پیش فاکتور"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            cell = itemtable.cell(len(items) + row_gap,2)
            cell.merge(itemtable.cell(len(items) + row_gap,1))
            cell = itemtable.cell(len(items) + row_gap,1)
            cell.paragraphs[0].text = formatNumber(str(total_items))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            cell = itemtable.cell(len(items) + row_gap,0)
            cell.paragraphs[0].text = formatNumber(str(total_price))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            ##################
            row_gap = 3
            itemtable.add_row()
            cell = itemtable.cell(len(items) + row_gap,4)
            cell.merge(itemtable.cell(len(items) + row_gap,3))
            cell = itemtable.cell(len(items) + row_gap,3)
            cell.merge(itemtable.cell(len(items) + row_gap,2))
            cell = itemtable.cell(len(items) + row_gap,2)
            cell.paragraphs[0].text = ": توضیحات"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(9)

            cell = itemtable.cell(len(items) + row_gap,1)
            cell.paragraphs[0].text = ": تخفیف"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            cell = itemtable.cell(len(items) + row_gap,0)
            cell.paragraphs[0].text = fa_number(metas['off']) + "%"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            total_price = round((total_price * (100-metas['off']))/100)

            ##################
            row_gap = 4
            itemtable.add_row()
            cell = itemtable.cell(len(items) + row_gap,4)
            cell.merge(itemtable.cell(len(items) + row_gap,3))
            cell = itemtable.cell(len(items) + row_gap,3)
            cell.merge(itemtable.cell(len(items) + row_gap,2))
            cell = itemtable.cell(len(items) + row_gap,2)
            cell.paragraphs[0].text = metas['text']
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            cell = itemtable.cell(len(items) + row_gap,1)
            cell.paragraphs[0].text = ": مالیات"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            cell = itemtable.cell(len(items) + row_gap,0)
            cell.paragraphs[0].text = formatNumber(metas['tax']) + " تومان"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(8)

            total_price += metas['tax']

            ##################
            row_gap = 5
            itemtable.add_row()
            cell = itemtable.cell(len(items) + row_gap,4)
            cell.paragraphs[0].text = "مبلغ قابل پرداخت"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(9)

            cell = itemtable.cell(len(items) + row_gap,3)
            cell.merge(itemtable.cell(len(items) + row_gap,2))
            cell = itemtable.cell(len(items) + row_gap,2)
            cell.merge(itemtable.cell(len(items) + row_gap,1))
            cell = itemtable.cell(len(items) + row_gap,1)
            cell.paragraphs[0].text = convert_to_words(total_price) + " تومان"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(9)

            cell = itemtable.cell(len(items) + row_gap,0)
            cell.paragraphs[0].text = formatNumber(str(total_price))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.paragraphs[0].style.font.size = Pt(9)

            ##################
            for x in doc.tables:
                x.style.font.name = "Vazir"

            ##################
            doc.save(output)

            ##################
            doc = docxtpl.DocxTemplate(output)
            doc.render(datas)
            doc.save(output)
            convert(output)

simpleFactor = {
    "server": {
        "name":"متانویا",
        "pid":"",
        "fid":"",
        "post":fa_number('1318917628'),
        "phone":fa_number('09966400214'),
        "loc":"",
    },
    "client": {
        "name":"",
        "pid":"",
        "fid":"",
        "post":"",
        "phone":"",
        "loc":"",
    },
    "item": {
        "0":["","",0,0],
    }
}

simpleFactor = {
    "write_date":fa_number('1385/02/24'),
    "factor_serial":fa_number('15248'),
    "addons":fa_number('1385/02/24'),

    "s_name":"متانویا",
    "s_id":fa_number("0770334891"),
    "s_post":fa_number("1318917628"),
    "s_phone":fa_number("09966400214"),
    "s_location":"",

    "c_name":"",
    "c_id":"",
    "c_post":"",
    "c_phone":"",
    "c_location":"",
}

items = [
    ["مبین","نمیدونم",5,2100],
    ["مانی","باید بدونم",4,5000],
    ["ارش","نخواستم",2,3400],
]

metas = {
    "off":10,
    "tax":1975,
    "text":"",
}

FactorMaker.createFactor("factors/template2.docx","factors/output.docx",simpleFactor,items,metas)