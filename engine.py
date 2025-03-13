import docx
from docx.enum.table import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import docxtpl
from docx2pdf import convert
import os

def argumentChecker(src_list:list,des_list:list):
    for x in des_list:
        if not x in src_list: return False
    return True

persian_numbers = {
    0: "صفر", 1: "یک", 2: "دو", 3: "سه", 4: "چهار", 5: "پنج", 6: "شش", 7: "هفت", 8: "هشت", 9: "نه",
    10: "ده", 11: "یازده", 12: "دوازده", 13: "سیزده", 14: "چهارده", 15: "پانزده", 16: "شانزده", 17: "هفده", 18: "هجده", 19: "نوزده",
    20: "بیست", 30: "سی", 40: "چهل", 50: "پنجاه", 60: "شصت", 70: "هفتاد", 80: "هشتاد", 90: "نود",
    100: "صد", 200: "دویست", 300: "سیصد", 400: "چهارصد", 500: "پانصد", 600: "ششصد", 700: "هفتصد", 800: "هشتصد", 900: "نهصد"
}

powers_of_thousand = [
    (1000000000, "میلیارد"),  # Billion
    (1000000, "میلیون"),      # Million
    (1000, "هزار"),           # Thousand
]

def_argument = {
    "makeFactor": {
        "datas": [
            "factor_id",
            "factor_date",
            "factor_addons",
            "title",'sby','phone','location',
            ''
        ]
    }
}

def_params = {
    "factor_id":"",
    "factor_date":"",
    "factor_addons":"",

    "title":"",
    "sby":"",
    "phone":"",
    "location":"",

    "factor_price_name":"",
    "factor_price_number":"",
    "factor_disc_1":"",
    "factor_disc_2":"",
}

image_map = {
    "blue":"template-1",
    # "cyan":"template-2",
    # "red":"template-3",
    # "yellow":"template-4",
    "gray":"template-gray",
}

class PersianFactor:
    def __init__(self):
        pass

    def makeFactor(self,ids:int,template_path:str,datas:dict):
        """ Make An Factor Using Templates """

        # output_path = f"output/{ids}.docx"

        # if not argumentChecker(list(datas.keys()),def_argument['makeFactor']['datas']): return False
        buffer = def_params.copy()
        buffer.update(datas)
        totalPrice = self.calculateTotal(list(buffer['items'].values()),datas['tax'],datas['off'])
        buffer['factor_price_name'] = self.convertWord(totalPrice) + " ریال" 
        buffer['factor_price_number'] = self.formatNumber(str(totalPrice))
        
        # print(buffer['factor_price_name'])
        # print(buffer['factor_price_number'])

        # PHASE 1
        doc = docx.Document(template_path)
        dtab = doc.tables[0]
        idx = 0
        for x in buffer['items']:
            buffer['items'][x]
            # if idx != 0: dtab.add_row()
            res = dtab.cell(idx,3)
            res.text = x
            res.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            res.paragraphs[0].style.font.size = Pt(7)
            res = dtab.cell(idx,2)
            res.text = self.persianNumber(str(buffer['items'][x][0]))
            res.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            res.paragraphs[0].style.font.size = Pt(7)
            res = dtab.cell(idx,1)
            res.text = self.formatNumber(str(buffer['items'][x][1]))
            res.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            res.paragraphs[0].style.font.size = Pt(7)
            res = dtab.cell(idx,0)
            res.text = self.formatNumber(str(round(buffer['items'][x][1] * buffer['items'][x][0])))
            res.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            res.paragraphs[0].style.font.size = Pt(7)
            idx += 1
        for x in doc.tables:
            x.style.font.name = "Vazir"

        for x in image_map:
            output_path = f"output/{ids}-{x}.docx"
            print(output_path)
            doc.save(output_path)
            doc = docxtpl.DocxTemplate(output_path)
            doc.replace_pic("background_image",f"images/{image_map[x]}.jpg")
            doc.render(buffer)
            doc.save(output_path)
            self.getPDF(output_path)


    def calculateTotal(self,items:list,tax:float,off:float):
        """ Calculate The Total Price In The END """
        buffer = items.copy()
        buff = 0
        for x in buffer:
            buff += x[0] * x[1]
        buff *= (100 + (tax))/100
        buff *= (100 - (off))/100
        return round(buff)

    def getPDF(self,path:str):
        """ Convert DocX Files To PDF """
        return convert(path)

    def convertHundred(self,number):
        """ Convert numbers from 1 to 999 to Persian words """
        if number < 20:
            return persian_numbers[number]
        elif number < 100:
            tens = number // 10 * 10
            remainder = number % 10
            return persian_numbers[tens] + ("" if remainder == 0 else " و " + persian_numbers[remainder])
        else:
            hundreds = number // 100 * 100
            remainder = number % 100
            return persian_numbers[hundreds] + ("" if remainder == 0 else " و " + self.convertHundred(remainder))

    def convertWord(self,number):
        """ Convert any number to Persian words """
        if number == 0:
            return persian_numbers[0]

        words = []

        # Iterate over powers of thousand
        for value, word in powers_of_thousand:
            if number >= value:
                count = number // value
                words.append(self.convertHundred(count) + " " + word)
                number %= value

        if number > 0:
            words.append(self.convertHundred(number))

        return " و ".join(words)

    def persianNumber(self,number:str) -> str:
        """ Convert Any English Number To Persian Format """
        perDigits = {
            '0': '۰', '1': '۱', '2': '۲', '3': '۳', '4': '۴', 
            '5': '۵', '6': '۶', '7': '۷', '8': '۸', '9': '۹'
        }
        return ''.join(perDigits[x] if x in perDigits else x for x in str(number))

    def formatNumber(self,number:str) -> str:
        """ Gives ، To Every 3 Numbers Arranged """
        output = str(self.persianNumber(number))[::-1]
        buffer = [output[x:x+3] for x in range(0,len(output),3)]
        return '،'.join(buffer)[::-1]

if __name__ == "__main__":
    d = PersianFactor()
    ids = 40312005
    # ids = 157565
    lts = {
        # "نصب و راه‌اندازی نرم‌افزار آریا":[11,50000],
        # "نصب و راه‌اندازی نرم‌افزار میشا و کوشا":[2,70000],
        # "فعال‌سازی و کرک نرم‌افزار آریا":[1,1300000],
        # "فعال‌سازی و کرک نرم‌افزار میشا و کوشا":[1,400000],
        # "بازیابی و ریکاوری سیستم":[1,250000],
        # "پاکسازی ویروس و بهینه‌سازی سیستم":[3,100000],
        # "فعال‌سازی پرینتر و راه‌اندازی اولیه":[1,100000]
        # "نصب آنتی ویروس":[1,1500000],
        "نصب و راه اندازی ویندوز":[2,1500000],
        "نصب درایور ها و برنامه ها":[2,1000000],
        "تعمیر و ارتقا سیستم":[2,7500000],
        "اسمبل سیستم":[2,2000000],
        "رم 4 گیگابایت DDR3":[2,5000000],
        # "موس":[1,3000000],
    }
    datas = {
        "factor_id":d.persianNumber(str(ids)),
        "factor_date":d.persianNumber("1403/11/15"),
        "factor_addons":"",

        "title":"فاکتور فروش",
        "sby":"تیم توسعه دهنده متانویا",
        "phone":"",
        "location":"تهران، منطقه 10، بلوار ستارخان، خیابان توکلی، انتهای کوچه خامنه، دبیرستان دخترانه ولایت فقیه",

        "items":lts,
        "off":0.0,
        "tax":0.0,

        "factor_disc_1":d.persianNumber("6219 - 8618 - 2621 - 0253\n6219 - 8619 - 8779 - 8203"),
        "factor_disc_2":"مبین براتیان تاج الدینی - مانی حامی کمری",
    }
    d.makeFactor(ids,"templates/template.docx",datas)